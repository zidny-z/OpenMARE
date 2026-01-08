import re
from pathlib import Path

from docx import Document
from docx.shared import Pt


def export_md_to_docx(md_text: str, path: str):
	doc = Document()
	doc.add_heading("Software Requirements Specification", level=1)

	lines = md_text.split("\n")
	in_code_block = False

	for line in lines:
		stripped = line.strip()

		# code
		if stripped.startswith("```"):
			in_code_block = not in_code_block
			continue

		if in_code_block:
			p = doc.add_paragraph(line)
			p.style.font.name = "Courier New"
			p.style.font.size = Pt(9)
			continue

		# h
		if stripped.startswith("### "):
			doc.add_heading(stripped[4:], level=3)
			continue

		if stripped.startswith("## "):
			doc.add_heading(stripped[3:], level=2)
			continue

		if stripped.startswith("# "):
			doc.add_heading(stripped[2:], level=1)
			continue

		# bulet
		if stripped.startswith(("- ", "* ")):
			p = doc.add_paragraph(stripped[2:], style="List Bullet")
			continue

		# number list
		if re.match(r"\d+\.\s+", stripped):
			p = doc.add_paragraph(stripped, style="List Number")
			continue

		# -
		if stripped == "":
			doc.add_paragraph("")
			continue

		# new line
		p = doc.add_paragraph()
		cursor = 0
		for match in re.finditer(r"(\*\*.*?\*\*|\*.*?\*)", line):
			p.add_run(line[cursor : match.start()])

			text = match.group()
			run = p.add_run(text.strip("*"))

			if text.startswith("**"):
				run.bold = True
			else:
				run.italic = True

			cursor = match.end()

		p.add_run(line[cursor:])

	Path(path).parent.mkdir(parents=True, exist_ok=True)
	doc.save(path)
